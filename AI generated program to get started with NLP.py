#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import pandas as pd
import nltk
from nltk.corpus import stopwords
from nltk.collocations import BigramCollocationFinder, TrigramCollocationFinder
from nltk.probability import FreqDist
from wordcloud import WordCloud
from textblob import TextBlob
from docx import Document
import matplotlib.pyplot as plt

# Ensure necessary NLTK data is downloaded
nltk.download('punkt')
nltk.download('stopwords')

# Load the survey comments from the Excel file
df = pd.read_excel('generated_survey_comments.xlsx')

# Tokenize the comments to identify keywords and remove stopwords
stop_words = set(stopwords.words('english'))
all_words = [word for word in nltk.word_tokenize(' '.join(df['Comment']).lower()) if word.isalnum() and word not in stop_words]
fdist = FreqDist(all_words)

# Define a function to create a theme dictionary based on the most common words
def create_theme_dict(fdist, num_keywords=50):
    common_words = fdist.most_common(num_keywords)
    theme_dict = {}
    for word, _ in common_words:
        theme_dict[word] = [word]
    return theme_dict

# Create the theme dictionary
theme_dict = create_theme_dict(fdist)

# Function to assign themes to comments
def assign_themes(comment, theme_dict):
    themes = []
    for theme, keywords in theme_dict.items():
        if any(keyword in comment.lower() for keyword in keywords):
            themes.append(theme)
    return themes

# Analyze sentiment and assign themes
df['Sentiment'] = df['Comment'].apply(lambda x: TextBlob(x).sentiment.polarity)
df['Themes'] = df['Comment'].apply(lambda x: assign_themes(x, theme_dict))

# Generate Word report
def generate_word_report(df):
    doc = Document()
    doc.add_heading('Survey Analysis Report', 0)
    doc.add_heading('Survey Question', level=1)
    doc.add_paragraph('What additional products or services would you like to see in the future at the campus bookstore?')

    # Overall sentiment score
    overall_sentiment = df['Sentiment'].mean()
    doc.add_heading('Overall Sentiment Score', level=1)
    doc.add_paragraph(f'{overall_sentiment:.2f}')

    # Collocations, bigrams, trigrams with number of comments for each
    words = [word for word in nltk.word_tokenize(' '.join(df['Comment']).lower()) if word.isalnum() and word not in stop_words]
    bigram_finder = BigramCollocationFinder.from_words(words)
    trigram_finder = TrigramCollocationFinder.from_words(words)
    bigrams = bigram_finder.ngram_fd.items()
    trigrams = trigram_finder.ngram_fd.items()

    doc.add_heading('Top 20 Collocations', level=1)
    doc.add_paragraph(', '.join([f'{bigram[0]} {bigram[1]} ({freq})' for bigram, freq in bigrams][:20]))
    doc.add_heading('Top 20 Bigrams', level=1)
    doc.add_paragraph(', '.join([f'{bigram[0]} {bigram[1]} ({freq})' for bigram, freq in bigrams][:20]))
    doc.add_heading('Top 10 Trigrams', level=1)
    doc.add_paragraph(', '.join([f'{trigram[0]} {trigram[1]} {trigram[2]} ({freq})' for trigram, freq in trigrams][:10]))

    # Wordcloud
    wordcloud = WordCloud(width=800, height=400).generate(' '.join(df['Comment']))
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.savefig('wordcloud.png')
    doc.add_heading('Wordcloud', level=1)
    doc.add_picture('wordcloud.png')

    # Sample comments with demographic information
    doc.add_heading('Sample Comments', level=1)
    for i, row in df.sample(5).iterrows():
        doc.add_paragraph(f"Comment: {row['Comment']}")
        doc.add_paragraph(f"Age: {row['Age']}, Gender: {row['Gender']}, Residency Campus: {row['Residency Campus']}, Indigenous: {row['Indigenous']}, Faculty: {row['Faculty']}, Program: {row['Program']}")

    doc.save('survey_analysis_report.docx')

# Generate Excel summary
def generate_excel_summary(df):
    theme_summary = df['Themes'].explode().value_counts().reset_index()
    theme_summary.columns = ['Theme', 'Count']
    theme_summary['Percentage'] = (theme_summary['Count'] / len(df)) * 100

    demographic_summary = df.groupby(['Residency Campus']).size().reset_index(name='Count')
    demographic_summary['Percentage'] = (demographic_summary['Count'] / len(df)) * 100

    summary_df = pd.merge(theme_summary, demographic_summary, how='outer', on=None)

    summary_df.to_excel('survey_summary.xlsx', index=False)

# Generate Excel file with comments and themes
def generate_comments_with_themes(df):
    df.to_excel('comments_with_themes.xlsx', index=False)

# Run the functions to generate outputs
generate_word_report(df)
generate_excel_summary(df)
generate_comments_with_themes(df)

print("The analysis has been completed and the files have been generated successfully.")


# In[ ]:




